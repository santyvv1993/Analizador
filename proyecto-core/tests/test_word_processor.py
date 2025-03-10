import pytest
from pathlib import Path
from datetime import datetime
import docx
from src.core.processors.word_processor import WordProcessor

# Definir ruta de recursos de prueba
TEST_RESOURCES = Path(__file__).parent / "resources"
TEST_RESOURCES.mkdir(exist_ok=True)

@pytest.fixture
def word_processor():
    return WordProcessor()

@pytest.fixture
def sample_docx_path():
    """Crea un documento Word de prueba"""
    docx_path = TEST_RESOURCES / "sample.docx"
    
    # Crear un documento nuevo
    doc = docx.Document()
    
    # Agregar título
    doc.add_heading('Test Document', 0)
    
    # Agregar párrafos de prueba
    doc.add_paragraph('This is a test document for Word processing.')
    doc.add_paragraph('It contains multiple paragraphs and formatting.')
    
    # Agregar una lista con viñetas
    doc.add_paragraph('Key points:', style='List Bullet')
    doc.add_paragraph('First important point', style='List Bullet')
    doc.add_paragraph('Second important point', style='List Bullet')
    
    # Agregar una tabla
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = 'Header 1'
    table.rows[0].cells[1].text = 'Header 2'
    table.rows[1].cells[0].text = 'Data 1'
    table.rows[1].cells[1].text = 'Data 2'
    
    # Establecer propiedades del documento
    doc.core_properties.author = "Test Author"
    doc.core_properties.title = "Test Document"
    doc.core_properties.subject = "Testing"
    doc.core_properties.category = "Test Documents"
    
    # Guardar el documento
    doc.save(docx_path)
    
    return str(docx_path)

@pytest.fixture
def sample_docx_path_with_content():
    """Crea un documento Word con contenido más significativo"""
    docx_path = TEST_RESOURCES / "sample_content.docx"
    
    doc = docx.Document()
    
    # Agregar encabezado
    doc.add_heading('Market Analysis Report', 0)
    
    # Agregar metadatos
    doc.core_properties.author = "Market Analyst"
    doc.core_properties.title = "Market Analysis Report 2024"
    doc.core_properties.category = "Business Reports"
    
    # Agregar contenido estructurado
    doc.add_paragraph('Date: February 15, 2024')
    doc.add_paragraph('This report analyzes technology market trends in 2024.')
    
    # Agregar sección de hallazgos
    doc.add_heading('Key Findings', level=1)
    findings = doc.add_paragraph()
    findings.add_run('Main discoveries include:')
    doc.add_paragraph('1. Increased demand for AI solutions', style='List Number')
    doc.add_paragraph('2. Growth in cloud computing market', style='List Number')
    doc.add_paragraph('3. New data privacy regulations', style='List Number')
    
    # Agregar tabla de empresas
    doc.add_heading('Companies Mentioned', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = 'Company'
    table.rows[0].cells[1].text = 'Sector'
    table.rows[1].cells[0].text = 'Microsoft Corporation'
    table.rows[1].cells[1].text = 'Technology'
    table.rows[2].cells[0].text = 'Amazon Web Services'
    table.rows[2].cells[1].text = 'Cloud Services'
    table.rows[3].cells[0].text = 'Google Cloud Platform'
    table.rows[3].cells[1].text = 'Cloud Services'
    
    # Agregar conclusiones
    doc.add_heading('Conclusions', level=1)
    doc.add_paragraph('The market shows sustained growth with emphasis on emerging technologies.')
    
    # Guardar el documento
    doc.save(docx_path)
    
    return str(docx_path)

def test_validate_with_valid_docx(word_processor, sample_docx_path):
    assert word_processor.validate(sample_docx_path) is True

def test_validate_with_invalid_path(word_processor):
    assert word_processor.validate("nonexistent.docx") is False

def test_get_mime_type(word_processor):
    assert word_processor.get_mime_type() == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

def test_process_valid_docx(word_processor, sample_docx_path):
    result = word_processor.process(sample_docx_path)
    
    # Verificar estructura básica
    assert isinstance(result.content, str)
    assert isinstance(result.metadata, dict)
    assert isinstance(result.summary, str)
    assert isinstance(result.keywords, list)
    assert isinstance(result.confidence_score, float)
    
    # Verificar metadatos específicos de Word
    assert result.author == "Test Author"
    assert result.title == "Test Document"
    assert result.metadata["subject"] == "Testing"
    assert result.metadata["category"] == "Test Documents"
    assert "document_statistics" in result.metadata

def test_process_invalid_docx(word_processor):
    with pytest.raises(ValueError):
        word_processor.process("nonexistent.docx")

def test_ai_analysis_integration(word_processor, sample_docx_path_with_content):
    """Prueba la integración con IA"""
    result = word_processor.process(sample_docx_path_with_content)
    
    # Verificar análisis de IA
    assert "ai_analysis" in result.metadata
    ai_analysis = result.metadata["ai_analysis"]
    
    # Verificar estructura básica
    assert isinstance(ai_analysis, dict)
    assert "success" in ai_analysis
    assert "analysis_result" in ai_analysis
    assert "confidence_score" in ai_analysis
    
    # Verificar contenido del análisis
    analysis_result = ai_analysis["analysis_result"]
    assert "summary" in analysis_result
    assert "keywords" in analysis_result
    assert "entities" in analysis_result

# Nuevas pruebas para funciones específicas del WordProcessor

def test_extract_tables(word_processor, sample_docx_path_with_content):
    """Prueba la extracción de tablas de un documento Word"""
    doc = docx.Document(sample_docx_path_with_content)
    tables_content = word_processor._extract_tables(doc)
    
    # Verificar que se extraen correctamente las tablas
    assert len(tables_content) > 0
    assert "Company | Sector" in tables_content[0]
    assert "Microsoft Corporation | Technology" in tables_content[0]
    assert "Amazon Web Services | Cloud Services" in tables_content[0]

def test_table_to_dict(word_processor, sample_docx_path_with_content):
    """Prueba la conversión de tablas a diccionario"""
    doc = docx.Document(sample_docx_path_with_content)
    table_dict = word_processor._table_to_dict(doc.tables[0])
    
    # Verificar estructura del diccionario
    assert "rows" in table_dict
    assert "dimensions" in table_dict
    assert table_dict["dimensions"]["rows"] == 4
    assert table_dict["dimensions"]["columns"] == 2
    
    # Verificar contenido de las filas
    assert table_dict["rows"][0][0] == "Company"
    assert table_dict["rows"][1][0] == "Microsoft Corporation"
    assert table_dict["rows"][1][1] == "Technology"

def test_extract_document_statistics(word_processor, sample_docx_path):
    """Prueba la extracción de estadísticas del documento"""
    doc = docx.Document(sample_docx_path)
    stats = word_processor._extract_document_statistics(doc)
    
    # Verificar estadísticas básicas
    assert "paragraphs" in stats
    assert "tables" in stats
    assert "sections" in stats
    assert "character_count" in stats
    assert "word_count" in stats
    
    # Verificar valores
    assert stats["paragraphs"] > 0
    assert stats["tables"] == 1
    assert stats["character_count"] > 0
    assert stats["word_count"] > 0

def test_count_pages(word_processor, sample_docx_path_with_content):
    """Prueba la estimación del número de páginas"""
    doc = docx.Document(sample_docx_path_with_content)
    pages = word_processor._count_pages(doc)
    
    # Verificar que se estima al menos 1 página
    assert pages >= 1
    
    # El documento de prueba debería tener al menos una página
    assert isinstance(pages, int)

def test_extract_keywords(word_processor):
    """Prueba la extracción de palabras clave"""
    content = """
    This is a test document for analyzing keyword extraction functionality.
    The document contains multiple occurrences of important keywords.
    Keywords should be extracted based on their frequency and relevance.
    Important keywords might include document, extraction, keywords, frequency, and relevance.
    """
    
    keywords = word_processor._extract_keywords(content)
    
    # Verificar estructura de las palabras clave
    assert isinstance(keywords, list)
    assert len(keywords) <= 10
    
    # Verificar palabras clave esperadas
    common_keywords = ["document", "keywords", "extraction", "frequency", "relevance"]
    found = [kw for kw in keywords if kw in common_keywords]
    assert len(found) > 0, "No se encontraron palabras clave esperadas"
