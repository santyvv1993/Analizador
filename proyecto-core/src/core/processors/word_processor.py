from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime
import docx
import langdetect
from .base_processor import BaseProcessor, ProcessedContent
from ..ai.ai_analyzer import AIAnalyzer

class WordProcessor(BaseProcessor):
    """Procesador específico para archivos Word (DOCX)"""

    def __init__(self):
        self.ai_analyzer = AIAnalyzer()

    def validate(self, file_path: str) -> bool:
        if not Path(file_path).exists():
            return False
        try:
            docx.Document(file_path)
            return True
        except Exception:
            return False

    def get_mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def process(self, file_path: str) -> ProcessedContent:
        if not self.validate(file_path):
            raise ValueError(f"Archivo inválido o no existe: {file_path}")

        doc = docx.Document(file_path)
        
        # Extraer contenido de párrafos
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)

        # Extraer tablas
        tables_data = self._extract_tables(doc)
        for table_content in tables_data:
            content.append(table_content)

        full_content = "\n".join(content)

        # Extraer metadatos
        core_properties = doc.core_properties
        document_stats = self._extract_document_statistics(doc)
        
        metadata = {
            "author": core_properties.author,
            "created": core_properties.created,
            "modified": core_properties.modified,
            "title": core_properties.title,
            "subject": core_properties.subject,
            "keywords": core_properties.keywords,
            "language": core_properties.language,
            "category": core_properties.category,
            "comments": core_properties.comments,
            "document_statistics": document_stats,
            "tables_count": len(doc.tables),
            "tables_data": [self._table_to_dict(table) for table in doc.tables]
        }

        # Realizar análisis con IA
        ai_analysis = self.ai_analyzer.analyze_content(full_content, metadata)
        analysis_result = ai_analysis.get("analysis_result", {})

        return ProcessedContent(
            content=full_content,
            metadata={
                **metadata,
                "ai_analysis": ai_analysis
            },
            summary=analysis_result.get("summary", full_content[:500]),
            keywords=analysis_result.get("keywords", self._extract_keywords(full_content)),
            created_date=core_properties.created,
            modified_date=core_properties.modified,
            author=core_properties.author,
            title=core_properties.title,
            num_pages=self._count_pages(doc),
            language=core_properties.language or langdetect.detect(full_content) if full_content else "en",
            entities=analysis_result.get("entities", []),
            confidence_score=ai_analysis.get("confidence_score", 0.5)
        )

    def _extract_tables(self, doc) -> List[str]:
        """Extrae el contenido de todas las tablas en el documento"""
        tables_content = []
        
        for table in doc.tables:
            table_text = []
            # Crear encabezados de tabla si es la primera fila
            header_row = []
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip())
            table_text.append(" | ".join(header_row))
            table_text.append("-" * len(" | ".join(header_row)))  # Separador
            
            # Resto de filas de la tabla
            for row in table.rows[1:]:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                table_text.append(" | ".join(row_text))
            
            tables_content.append("\n".join(table_text))
        
        return tables_content
    
    def _table_to_dict(self, table) -> Dict:
        """Convierte una tabla a formato diccionario para metadata"""
        result = {"rows": []}
        
        for i, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            result["rows"].append(row_data)
        
        result["dimensions"] = {
            "rows": len(table.rows),
            "columns": len(table.columns) if table.rows else 0
        }
        
        return result

    def _extract_document_statistics(self, doc) -> Dict:
        """Extrae estadísticas del documento"""
        stats = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
            "character_count": 0,
            "word_count": 0
        }
        
        # Contar palabras y caracteres
        for para in doc.paragraphs:
            text = para.text.strip()
            stats["character_count"] += len(text)
            stats["word_count"] += len(text.split())
        
        return stats

    def _count_pages(self, doc) -> int:
        """
        Aproximación del número de páginas.
        La API de python-docx no proporciona acceso directo al conteo de páginas,
        así que usamos el número de secciones como aproximación.
        """
        # Una mejor aproximación sería:
        # 1 página ≈ 3000 caracteres o 500 palabras
        char_count = sum(len(para.text) for para in doc.paragraphs)
        word_count = sum(len(para.text.split()) for para in doc.paragraphs)
        
        # Usar el mayor de los dos estimados
        pages_by_char = max(1, char_count // 3000)
        pages_by_word = max(1, word_count // 500)
        
        return max(pages_by_char, pages_by_word, len(doc.sections))

    def _extract_keywords(self, content: str) -> List[str]:
        """Extrae palabras clave del contenido basado en frecuencia"""
        # Eliminar palabras comunes y quedarse con palabras relevantes
        words = content.lower().split()
        keywords = [word for word in words if len(word) > 4]
        
        from collections import Counter
        return [word for word, _ in Counter(keywords).most_common(10)]
